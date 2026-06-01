from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUT_DIR = Path(r"C:\Users\KIM-LEGION-5-2025\Documents\Codex\GIU Tracker\GIU Tracker Other Files")
PORTFOLIO_OUT = OUT_DIR / "Project Portfolio Notes - GIU NRL Status Tracker - Updated.docx"
ROADMAP_OUT = OUT_DIR / "Roadmap - GIU NRL Status Tracker - Updated.docx"

ACCENT = RGBColor(23, 73, 90)
TEAL = RGBColor(11, 117, 111)
MUTED = RGBColor(96, 112, 124)
LIGHT_FILL = "EEF5F5"
BLUE_FILL = "E8F2FF"
GOLD_FILL = "FFF4C7"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_text(cell, text, bold=False, color=None):
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run(text)
    run.bold = bold
    run.font.name = "Aptos"
    run.font.size = Pt(9.5)
    if color:
        run.font.color.rgb = color


def set_borders(table):
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = f"w:{edge}"
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), "4")
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), "D7E2E8")


def setup_document(title, subtitle):
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.75)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(0.85)
    section.right_margin = Inches(0.85)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Aptos"
    normal.font.size = Pt(10.5)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.12

    for style_name, size, color, before, after in [
        ("Heading 1", 15, ACCENT, 12, 5),
        ("Heading 2", 12.5, TEAL, 8, 4),
        ("Heading 3", 11, ACCENT, 6, 3),
    ]:
        style = styles[style_name]
        style.font.name = "Aptos Display"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = color
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)

    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    title_run = title_para.add_run(title)
    title_run.font.name = "Aptos Display"
    title_run.font.size = Pt(20)
    title_run.font.bold = True
    title_run.font.color.rgb = ACCENT
    title_para.paragraph_format.space_after = Pt(2)

    subtitle_para = doc.add_paragraph()
    subtitle_run = subtitle_para.add_run(subtitle)
    subtitle_run.font.name = "Aptos"
    subtitle_run.font.size = Pt(10.5)
    subtitle_run.font.color.rgb = MUTED
    subtitle_para.paragraph_format.space_after = Pt(10)

    add_rule(doc)
    return doc


def add_rule(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(8)
    p_pr = p._p.get_or_add_pPr()
    p_bdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "8")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "D7E2E8")
    p_bdr.append(bottom)
    p_pr.append(p_bdr)


def para(doc, text):
    p = doc.add_paragraph(text)
    p.paragraph_format.space_after = Pt(6)
    return p


def bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(2)
        p.add_run(item)


def numbered(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Number")
        p.paragraph_format.space_after = Pt(2)
        p.add_run(item)


def key_value_table(doc, rows):
    table = doc.add_table(rows=1, cols=2)
    table.autofit = False
    table.columns[0].width = Inches(2.0)
    table.columns[1].width = Inches(4.6)
    set_borders(table)
    for i, (key, value) in enumerate(rows):
        cells = table.rows[0].cells if i == 0 else table.add_row().cells
        set_cell_text(cells[0], key, bold=True, color=ACCENT)
        set_cell_text(cells[1], value)
        set_cell_shading(cells[0], LIGHT_FILL)
    doc.add_paragraph()
    return table


def status_table(doc, headers, rows):
    table = doc.add_table(rows=1, cols=len(headers))
    table.autofit = True
    set_borders(table)
    for cell, header in zip(table.rows[0].cells, headers):
        set_cell_text(cell, header, bold=True, color=ACCENT)
        set_cell_shading(cell, LIGHT_FILL)
    for row in rows:
        cells = table.add_row().cells
        for cell, text in zip(cells, row):
            set_cell_text(cell, text)
    doc.add_paragraph()
    return table


def footer(doc):
    section = doc.sections[0]
    footer_p = section.footer.paragraphs[0]
    footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer_p.add_run("GIU NRL Status Tracker | Updated May 31, 2026 | Prepared by Kim Pusing")
    run.font.name = "Aptos"
    run.font.size = Pt(8.5)
    run.font.color.rgb = MUTED


def build_portfolio():
    doc = setup_document(
        "GIU NRL Status Tracker - Project Portfolio Notes",
        "Updated portfolio summary for the current requester build and pilot prototype.",
    )

    doc.add_heading("Project Information", level=1)
    key_value_table(
        doc,
        [
            ("Project name", "GIU NRL Status Tracker"),
            ("Project type", "Internal workflow management and request tracking prototype"),
            ("Organization", "RITM Department of Virology - Genomics and Innovation Unit"),
            ("Current phase", "Pilot testing and user experience refinement"),
            ("Current build focus", "Requester details, request summaries, sign-in experience, pagination, and frontend usability polish"),
            ("Development period", "May 25-31, 2026"),
        ],
    )

    doc.add_heading("Project Summary", level=1)
    para(
        doc,
        "The GIU NRL Status Tracker is a web-based prototype designed to improve visibility, coordination, and status monitoring for GIU-supported requests from National Reference Laboratories under the Department of Virology.",
    )
    para(
        doc,
        "The prototype provides a shared dashboard where GIU staff can create and update request records, while approved users can view request progress, request summaries, status notes, next steps, and key timeline information.",
    )
    para(
        doc,
        "The project continues to serve as both a practical internal tool and a learning portfolio project covering AI-assisted development, frontend design, Supabase integration, role-based access, deployment workflows, and product iteration based on stakeholder feedback.",
    )

    doc.add_heading("Problem Statement", level=1)
    bullets(
        doc,
        [
            "GIU request updates were previously spread across email, chat messages, and direct follow-ups.",
            "Repeated status inquiries made it harder to maintain a single source of truth.",
            "GIU workload, next steps, and request progress were difficult to review at a glance.",
            "NRL stakeholders needed a clearer way to view relevant request status without relying only on manual follow-up.",
        ],
    )

    doc.add_heading("Solution", level=1)
    bullets(
        doc,
        [
            "A centralized web dashboard for tracking active, on-hold, delayed, and completed GIU work requests.",
            "Request creation and update tools for GIU/admin users.",
            "A clickable work queue with full request summaries for faster context review.",
            "Role-based access with Admin, GIU, NRL, and Pending User profiles.",
            "Supabase-backed persistence, authentication, and user profile management.",
        ],
    )

    doc.add_heading("Technology Stack", level=1)
    status_table(
        doc,
        ["Area", "Technology"],
        [
            ("Frontend", "HTML, CSS, JavaScript"),
            ("Backend and database", "Supabase with PostgreSQL"),
            ("Authentication", "Supabase Authentication"),
            ("Hosting", "Render"),
            ("Version control", "GitHub"),
            ("Development approach", "AI-assisted development with Codex and ChatGPT"),
        ],
    )

    doc.add_heading("Major Completed Updates", level=1)
    doc.add_heading("Request saving and creation", level=2)
    bullets(
        doc,
        [
            "Repaired the issue where creating a new request could fail or leave the modal stuck.",
            "Added clear save success and save error feedback inside the request modal.",
            "Improved duplicate request ID handling and save-state feedback.",
        ],
    )
    doc.add_heading("Requester and request summary module", level=2)
    bullets(
        doc,
        [
            "Integrated requester name, requester NRL/section, and required assistance into saved request records.",
            "Added clickable request summaries showing full request context.",
            "Included requester details, assigned staff, program, dates, GIU stage, status notes, and next step in the summary view.",
        ],
    )
    doc.add_heading("Work queue usability", level=2)
    bullets(
        doc,
        [
            "Improved queue row spacing and visual grouping for Program, GIU Stage, dates, and action buttons.",
            "Added pagination to the Queue view to reduce long scrolling.",
            "Added pagination to the By NRL view for easier browsing when requests increase.",
            "Maintained search and filters for laboratory, GIU stage, status, and priority.",
        ],
    )
    doc.add_heading("Sign-in and account experience", level=2)
    bullets(
        doc,
        [
            "Replaced the blank signed-out dashboard with a proper sign-in landing screen.",
            "Added friendly login error feedback for incorrect email or password.",
            "Improved create-account feedback for duplicate or failed signups.",
            "Added a Forgot Password option and password reset email confirmation message.",
            "Added Enter-key submission behavior for sign-in and create-account forms.",
        ],
    )
    doc.add_heading("Visual polish and branding", level=2)
    bullets(
        doc,
        [
            "Improved header action spacing for Admin view, Sign out, Manage Users, and New Request.",
            "Added the RITM logo as the browser tab icon.",
            "Refined sign-in page layout, accent treatment, and footer spacing.",
            "Improved responsiveness and overall prototype polish.",
        ],
    )

    doc.add_heading("Current Known Limitation", level=1)
    para(
        doc,
        "The Forgot Password email is now being sent successfully, but the reset link currently returns users to the site without presenting a dedicated form and logic to set a new password. This is planned as the next frontend task.",
    )

    doc.add_heading("Backend and Security Notes", level=1)
    bullets(
        doc,
        [
            "The Supabase database is currently aligned with the frontend fields used by the prototype.",
            "Future backend hardening should tighten pending-user row-level security so users cannot self-assign a real laboratory before approval.",
            "Function definitions, RLS-enabled status, triggers, and execution privileges should be verified before wider deployment.",
            "Audit logging, archive/soft delete, backup planning, and recovery procedures remain important future work.",
        ],
    )

    doc.add_heading("Lessons Learned", level=1)
    bullets(
        doc,
        [
            "End-to-end web application development from concept to deployed prototype.",
            "Supabase database integration, authentication, and role-based access patterns.",
            "Iterative UI/UX improvement based on real use and stakeholder feedback.",
            "Importance of separating frontend deployment from backend schema and security changes.",
            "Practical use of AI-assisted development for debugging, design iteration, and documentation.",
        ],
    )

    doc.add_heading("Portfolio Reflection", level=1)
    para(
        doc,
        "This project began as a focused learning exercise and has grown into a functional pilot prototype for internal workflow visibility. The current build demonstrates request tracking, user access control, Supabase-backed persistence, deployment through Render, and iterative product improvements based on practical testing.",
    )
    para(
        doc,
        "The next stage is to continue refining user experience, complete the password reset flow, make dashboard cards actionable, and then reinforce backend security before broader use.",
    )

    footer(doc)
    doc.save(PORTFOLIO_OUT)


def build_roadmap():
    doc = setup_document(
        "GIU NRL Status Tracker - Development Roadmap",
        "Updated roadmap after requester module, authentication UX, queue summary, and pagination improvements.",
    )

    doc.add_heading("Current Project Status", level=1)
    key_value_table(
        doc,
        [
            ("Current phase", "Pilot testing and frontend usability refinement"),
            ("Current build", "Current Requester Build"),
            ("Prototype readiness", "Usable for controlled prototype testing with approved users"),
            ("Primary focus", "Improve user experience and feedback before deeper backend hardening"),
        ],
    )

    doc.add_heading("Completed Since Initial Pilot Baseline", level=1)
    status_table(
        doc,
        ["Area", "Completed update"],
        [
            ("Request save flow", "Fixed create/save issue and added save success/error feedback."),
            ("Requester details", "Added requester name, requester NRL/section, and required assistance to saved request context."),
            ("Request summaries", "Made queue items clickable and added full request summary view."),
            ("Queue layout", "Improved row spacing and visual organization."),
            ("Pagination", "Added pagination for Queue and By NRL views."),
            ("Sign-in UX", "Added proper signed-out landing page, login errors, account creation feedback, and Enter-key submit."),
            ("Password reset email", "Added forgot-password action and reset email sent confirmation."),
            ("Branding/UI polish", "Added RITM favicon, improved header buttons, sign-in page, and footer spacing."),
        ],
    )

    doc.add_heading("Immediate Next Tasks", level=1)
    doc.add_heading("1. Complete password reset flow", level=2)
    bullets(
        doc,
        [
            "Detect Supabase password recovery sessions after users click the reset email link.",
            "Show a dedicated Set New Password form in the app.",
            "Update the password through Supabase and show clear success/error feedback.",
            "Redirect users back to the normal sign-in or dashboard flow after completion.",
        ],
    )
    doc.add_heading("2. Make dashboard cards clickable", level=2)
    bullets(
        doc,
        [
            "Make Active Requests, On Hold, Completed, and similar cards filter the queue when clicked.",
            "Add visual feedback so users understand which dashboard shortcut is active.",
            "Include a simple way to clear the dashboard shortcut filter.",
        ],
    )
    doc.add_heading("3. Continue frontend feedback improvements", level=2)
    bullets(
        doc,
        [
            "Refine empty states, error messages, and confirmation dialogs.",
            "Improve manage-user messaging and admin safeguards.",
            "Review mobile and smaller-screen behavior after pagination and layout changes.",
        ],
    )

    doc.add_heading("Backend Hardening Roadmap", level=1)
    status_table(
        doc,
        ["Priority", "Task", "Purpose"],
        [
            ("High", "Tighten pending-user RLS", "Prevent pending users from self-assigning a real lab and seeing request data before approval."),
            ("High", "Verify RLS is enabled", "Confirm RLS is active on profiles, giu_requests, and related tables as appropriate."),
            ("High", "Review SQL functions", "Confirm current_profile_role, current_profile_lab, reserve_giu_display_id, and touch_updated_at are safe and correctly scoped."),
            ("High", "Audit display ID reservation", "Ensure request ID allocation remains atomic under concurrent inserts."),
            ("Medium", "Limit function privileges", "Restrict function execution to intended authenticated roles where appropriate."),
            ("Medium", "Attach/verify updated_at trigger", "Ensure updated_at changes reliably for sorting and audit support."),
        ],
    )

    doc.add_heading("Data Protection and Governance", level=1)
    bullets(
        doc,
        [
            "Implement audit logging for request updates, user changes, status changes, and deletes.",
            "Replace hard delete with archive or soft-delete where appropriate.",
            "Create backup and restore procedures for Supabase data.",
            "Document emergency recovery steps for admin access and user profile issues.",
            "Create a repeatable migration process so database changes are tracked alongside GitHub code changes.",
        ],
    )

    doc.add_heading("Reporting and Analytics Opportunities", level=1)
    bullets(
        doc,
        [
            "Turnaround time summaries.",
            "Request volume by NRL, program, status, priority, and GIU staff member.",
            "Monthly accomplishment reporting support.",
            "Export options such as Excel, CSV, or PDF.",
            "Dashboard trends for active, delayed, completed, and on-hold requests.",
        ],
    )

    doc.add_heading("Hosting and Production Readiness", level=1)
    bullets(
        doc,
        [
            "Review Render free-tier limitations, including possible cold starts, loading delays, and performance constraints as user activity grows.",
            "Evaluate whether a paid hosting tier or alternative hosting setup is needed before wider deployment.",
            "Set up environment documentation for frontend deployment, Supabase configuration, and email settings.",
            "Improve Supabase email branding further after an official GIU email or SMTP provider is available.",
            "Prepare a short user guide for GIU, admin, and NRL users.",
        ],
    )

    doc.add_heading("Suggested Phase Plan", level=1)
    numbered(
        doc,
        [
            "Finish frontend user experience items: password reset form, clickable dashboard cards, and small usability refinements.",
            "Conduct another round of GIU/admin review using screenshots and prototype walkthrough.",
            "Implement backend security hardening and migration tracking.",
            "Add audit logging, archive/soft-delete, and backup procedures.",
            "Review hosting limits and prepare production-readiness documentation.",
            "Expand reporting and analytics features if the prototype continues to be useful for team workflows.",
        ],
    )

    doc.add_heading("Project Note", level=1)
    para(
        doc,
        "The prototype is now substantially more usable than the initial pilot baseline, especially for request intake, request review, and signed-out user experience. It should still be treated as a controlled prototype until the password reset completion flow, backend security hardening, audit logging, and backup procedures are completed.",
    )

    footer(doc)
    doc.save(ROADMAP_OUT)


if __name__ == "__main__":
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    build_portfolio()
    build_roadmap()
    print(PORTFOLIO_OUT)
    print(ROADMAP_OUT)
